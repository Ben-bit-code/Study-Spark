import sys, os, re, time
from llama_cpp import Llama
from PySide6.QtWidgets import QMessageBox, QWidget, QApplication, QFileDialog
from PySide6.QtCore import QThread, Signal
from ui_form import Ui_Widget
from textwrap import wrap
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
import pydot

#Worker Thread to prevent our GUI from freezing
class WorkerThread(QThread):
    result_ready = Signal(str)
    #The status variable will be sent to the main GUI to tell our main GUI the status of our programming
    status = Signal(str)
    def __init__(self, text, noteTakingMethod, path):
        super().__init__()
        self.text = text
        self.noteTakingMethod = noteTakingMethod
        self.path = path

    def run(self):
        try:
            #Emitting the status of our program to the main GUI
            self.status.emit("Loading Model")
            #Pausd to prevent signalling overload
            time.sleep(0.5)
            #Loading our model
            llm = Llama(
                model_path="models/Phi-3-mini-4k-instruct-q4.gguf", #Our model
                # seed=1337, # Uncomment to set a specific seed
                n_ctx=2048, # Max context window
            )
            self.status.emit("Wrapping Data")
            time.sleep(0.5)
            #Breaking up our text into a list
            finalData = wrap(self.text, 3600)
            notes = ""
            self.status.emit("Prompting Model")
            time.sleep(0.5)
            '''
            Running each element of our finalData element as a different element since our model can only take 2048
            tokens or 8192 characters
            '''
            for i in range(len(finalData)):
                if self.noteTakingMethod == "Outline Method":
                    output = llm(
                        f"""
                        Q: Summarize the following text using the Outline Method.
                        Use ALL CAPS for main ideas, bullet points (•) for subpoints,
                        hyphens (-) for further detail, and lowercase letters (a, b, c) for examples.


                        Text:
                        {finalData[i]}

                        Outline:
                        """,
                        max_tokens=768, # Maximum amount of creativity
                        temperature=0.3, # Our models allowed creativity (between 0 and 1)
                        top_p=0.9, # Controls diversity and randomness
                        repeat_penalty=1.1, # Prevents repition of information
                        stop=["Q:", "Question", "Answer", "End", "Response", "- ["], # Tells our model where to stop generating
                        echo=False) # Prevents the prompt from being repeated in the response
                    # Adding our required text to the notes variable
                    notes += output["choices"][0]["text"]
                elif self.noteTakingMethod == "Cornell Method":
                    output = llm(
                        f"""
                        Q: Take structured notes using the Cornell Method.
                        Create three sections: Cues (key questions or terms), Notes (main ideas),
                        and Summary (brief conclusion).

                        Text:
                        {finalData[i]}

                        A:
                        """,
                        max_tokens=640,
                        temperature=0.3,
                        top_p=0.85,
                        repeat_penalty=1.05,
                        stop=["Q:", "Question", "Answer", "End", "Response", "###", "<", "- ["],
                        echo=False)
                    notes += output["choices"][0]["text"]
                elif self.noteTakingMethod == "Boxing Method":
                    output = llm(
                        f"""
                        Q: Organize the following information using the Boxing Method.

                        Group related concepts into clearly separated sections (boxes),
                        each with a label and a short 1–2 sentence explanation.

                        ⚠️ Make sure the **pipe (|)** character appears **at the very end of the explanation**, not after the heading.

                        Repeat: the **| should be at the end of the explanation sentence, not after the title or anywhere else.**
                        Do not add the pipe after the heading/title.

                        Text:
                        {finalData[i]}

                        A:
                        """,
                        max_tokens=640,
                        temperature=0.3,
                        top_p=0.85,
                        repeat_penalty=1.05,
                        stop=["Q:", "Question", "Answer", "End", "Response", "###", "<", "- ["],
                        echo=False)
                    notes += output["choices"][0]["text"]
                    notes = notes.replace("_", "")
                elif self.noteTakingMethod == "Charting Method":
                    output = llm(
                        f"""
                        Q: Summarize the following content using the Charting Method. Use a three-column format:
                        Topic | Definition | Example |

                        Each row must be separated by a newline, and each column by a pipe symbol. Keep entries short, clear, and informative.


                        Text:
                        {finalData[i]}

                        A:
                        """,
                        max_tokens=512,
                        temperature=0.3,
                        top_p=0.85,
                        repeat_penalty=1.05,
                        stop=["Q:", "Question", "Answer", "End", "Response", "###", "<","- ["],
                        echo=False)
                    notes += output["choices"][0]["text"]
                elif self.noteTakingMethod == "Mapping Method":
                    output = llm(
                        f"""
                        You are a helpful assistant that creates structured mind maps from educational or technical content.

                        Given a topic, create a mind map with:
                        - **One single root node** at the top.
                        - Clear top-down structure (root → subtopics → details).
                        Format:

                        |Root Node| -> |Subtopic 1|;
                        |Root Node| -> |Subtopic 2|;
                        |Subtopic 1| -> |Detail A|;
                        |Subtopic 1| -> |Detail B|;
                        |Subtopic 2| -> |Detail C|;

                        Limit each detail and subtopic to 30 characters
                        Only display the mind map. Do not display the author.
                        Add a percentage (%) character at the end of the text.
                        Text:
                        {finalData[i]}
                        """,
                        max_tokens=768,
                        temperature=0.3,
                        top_p=0.85,
                        repeat_penalty=1.05,
                        stop=["Q:", "Question", "Answer", "End", "Response", "###", "<", "- [", "Written by"],
                        echo=False)
                    notes += output["choices"][0]["text"]
            self.status.emit("Creating Document")
            time.sleep(0.5)
            # Creating our document object
            document = Document()
            if self.noteTakingMethod == "Outline Method":
                document.add_paragraph(notes)
                document.save(self.path)
            elif self.noteTakingMethod == "Cornell Method":
                notes = re.sub(r'\s+',' ',notes)
                notes = re.split(r'(\d+)', notes)
                notes = tuple(notes)
                # Breaking up our data into a list that can be used to create a table with the python-docx library
                notes = tuple(notes[i:i + 7] for i in range(0, len(notes), 7))
                # Creating a table to neatly store our data
                table = document.add_table(rows=2, cols=2)
                table.style = 'ColorfulShading-Accent1'
                for i, num1, cue, num2, notes, num3, summary in notes:
                    # First row to store cues on the left and notes on the right
                    row_cells = table.add_row().cells
                    row_cells[0].text = cue
                    row_cells[1].text = notes
                    # Second merged row to store the summary of the text
                    row_cells = table.add_row().cells
                    row_cells[0].merge(row_cells[1])
                    row_cells[0].text = summary
                document.save(self.path)
            elif self.noteTakingMethod == "Boxing Method":
                notes = re.sub(r"(\d+\.\s.+?) \| (.+)", r"\1\n\2 |", notes)
                notes = notes.split("|")
                table = document.add_table(rows=1, cols=1)
                for box in notes:
                    row_cells = table.add_row().cells
                    row_cells[0].text = box
                document.save(self.path)
            elif self.noteTakingMethod == "Charting Method":
                notes = re.sub(r'\s+',' ',notes)
                notes = notes.split("|")
                del notes[-1]
                notes = tuple(notes)
                # We repeat some of what was performed in the Cornell Method
                notes = tuple(notes[i:i + 3] for i in range(0, len(notes), 3))
                table = document.add_table(rows=1, cols=3)
                table.style = 'ColorfulShading-Accent1'
                for topic, definition, example in notes:
                    row_cells = table.add_row().cells
                    row_cells[0].text = topic
                    row_cells[1].text = definition
                    row_cells[2].text = example
                document.save(self.path)
            elif self.noteTakingMethod == "Mapping Method":
                notes = notes.replace("Written by:", "")
                notes = notes.replace(".", "")
                notes = re.sub(r'\s+',' ',notes)
                # Breaking up our text into a list
                notes = notes.split("%")
                for i in range(len(notes)):
                    # Using our response to create dot string for our graph
                    dot_string = f""" graph mind_map {
                        {notes[i]}
                    }
                    """
                    # Removing unnecessary characters
                    dot_string = dot_string.replace("|", "\"")
                    dot_string = dot_string.replace("\'", "")
                    graphs = pydot.graph_from_dot_data(dot_string)
                    #Creating our graph
                    graph = graphs[0]
                    graph.write_png("output.png")
                    # Turning a regular word document page landscape
                    section = document.sections[-1]
                    new_width, new_height = section.page_height, section.page_width
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = new_width
                    section.page_height = new_height
                    # Adding our image to the word document
                    document.add_picture("output.png", width=Inches(9), height=Inches(3))
                    # Removing the image using the builtin os module
                    os.remove("output.png")
                # Saving our document
                document.save(self.path)
            self.status.emit("Saved")
        except Exception as e:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.setText(f"Error: {e}")
            msg.setWindowTitle('Error')
            msg.setStandardButtons(QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
            msg.exec()

#Our main GUI
class Widget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Widget()
        self.ui.setupUi(self)
        self.ui.reset.clicked.connect(self.resetData)
        self.ui.submit.clicked.connect(self.submitData)
        self.setWindowTitle("Study Spark")

    def saveFile(self):
        # Function only saves as docx file and returns the name and location as a string
        filename, filter = QFileDialog.getSaveFileName(parent=self, caption='Save file', dir='.', filter='Word Document (*.docx)')
        if filename:
            return str(filename)

    def resetData(self):
        # Resets all radiobuttons and textbox
        radioButtons = [self.ui.boxingMethodButton, self.ui.chartingMethodButton, self.ui.cornellMethodButton, self.ui.mappingMethodButton, self.ui.outlineMethodButton]
        for radioButton in radioButtons:
            radioButton.setAutoExclusive(False)
            radioButton.setChecked(False)
            radioButton.setAutoExclusive(True)
        self.ui.textInput.clear()
        self.ui.progressBar.reset()
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Icon.Information)
        msg.setText("Form data reset successfully")
        msg.setWindowTitle("Information")
        msg.setStandardButtons(QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
        msg.exec()

    def currentStatus(self, status):
        # The status is passed from the worker thread to this function to modify the progress bar
        if status == "Loading Model":
            for i in range(1, 11):
                self.ui.progressBar.setValue(i)
                time.sleep(0.05)
        if status == "Wrapping Data":
            for i in range(11, 16):
                self.ui.progressBar.setValue(i)
                time.sleep(0.05)
        if status == "Prompting Model":
            for i in range(16, 26):
                self.ui.progressBar.setValue(i)
                time.sleep(0.05)
        if status == "Creating Document":
            for i in range(26, 76):
                self.ui.progressBar.setValue(i)
                time.sleep(0.05)
        if status == "Saved":
            for i in range(76, 101):
                self.ui.progressBar.setValue(i)
                time.sleep(0.05)
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Information)
            msg.setText("Document saved successfully")
            msg.setWindowTitle("Notification")
            msg.setStandardButtons(QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
            msg.exec()

    def submitData(self):
        if self.ui.textInput.toPlainText() == "":
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setText("Fill in the text area provided")
            msg.setWindowTitle("Warning")
            msg.setStandardButtons(QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
            msg.exec()
            return False
        noteTakingMethod = ""
        if self.ui.outlineMethodButton.isChecked():
            noteTakingMethod = "Outline Method"
        elif self.ui.cornellMethodButton.isChecked():
            noteTakingMethod = "Cornell Method"
        elif self.ui.boxingMethodButton.isChecked():
            noteTakingMethod = "Boxing Method"
        elif self.ui.chartingMethodButton.isChecked():
            noteTakingMethod = "Charting Method"
        elif self.ui.mappingMethodButton.isChecked():
            noteTakingMethod = "Mapping Method"
        else:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setText("Please select a note taking method")
            msg.setWindowTitle("Warning")
            msg.setStandardButtons(QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
            msg.exec()
            return False

        path = self.saveFile()

        if path:
            #Creating an instance og our worker thread class
            self.workerThread = WorkerThread(self.ui.textInput.toPlainText(), noteTakingMethod, path)
            #Getting the status from our worker thread
            self.workerThread.status.connect(self.currentStatus)
            #starting our worker thread
            self.workerThread.start()
        else:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setText("Please select a valid path")
            msg.setWindowTitle("Warning")
            msg.setStandardButtons(QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
            msg.exec()
            return False

if __name__ == "__main__":
    app = QApplication(sys.argv)
    widget = Widget()
    widget.show()
    sys.exit(app.exec())
